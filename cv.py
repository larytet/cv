# Usage
# python3 -m venv .venv
# source .venv/bin/activate
# pip install -r requirements.txt
# sudo apt-get update
# sudo apt-get install -y pandoc texlive-latex-base texlive-latex-recommended texlive-latex-extra
# python3 cv.py

from docx import Document
import pypandoc
from docx2pdf import convert

ADD_TECHOLOGIES = True
IS_BACKEND = False
LONG_VERSION = False

# Create a new Document with the updated order
doc = Document()

# Add the content to the Document
doc.add_heading('Arkady Miasnikov', 0)
doc.add_paragraph('(054) 4681517 | arkady.miasnikov@gmail.com')

doc.add_heading('Software Engineer experienced in leading development projects from inception to full implementation', level=1)
if IS_BACKEND:
    doc.add_paragraph(
        "Backend and distributed-systems engineer with hands-on ownership of trading, security, and data-intensive platforms. "
        "Strong track record building performance-critical services in Go/Python, optimizing hot paths, and shipping resilient production systems. "
        "Experienced in cloud-native operations, observability, and cross-functional delivery from design through rollout. "
        "Tools & Technologies: Linux, Go, Python, C, AWS, Kubernetes, SQL, ClickHouse, Redis, Kafka."
    )
else:
    doc.add_paragraph(
        "Real-time and embedded software engineer with deep experience across firmware, Linux kernel, BSP/driver development, and hardware-software integration. "
        "Built and optimized latency-sensitive, multi-threaded systems from bring-up to production support on diverse architectures. "
        "Strong background in reverse engineering, low-level debugging, and system performance analysis across Linux and Windows kernel environments. "
        "Tools & Technologies: Linux kernel, C/C++, Python, Go, Assembly, RTOS, Yocto, U-Boot, eBPF/SystemTap."
    )

doc.add_heading('Professional Experience', level=1)

# Add the remaining experience
experience = [
    ("2023-present: Software Engineer - Endotech", [
        "Designed and implemented custom momentum indicators and market-condition filters used in production trading strategies.",
        "Ported and optimized multiple indicators from PineScript to Go to improve runtime efficiency and maintainability.",
        "Delivered reusable signal components that became core building blocks in the company's algorithmic trading stack."
    ], ['Go', 'Python', "ClickHouse", "SQL", "AI"], False),
    ("2022-2023: Software Engineer - HiAuto", [
        "Maintained and improved a computerized drive-thru order-taking platform in a distributed international engineering team.",
        "Integrated operational data into analytics workflows and resolved data quality/system issues to improve business visibility.",
        "Supported continuous operation of on-premises devices through proactive troubleshooting and incident resolution."
    ], ['Python', "BigQuery", "Redis", "SQL", "AI", "k8s", "Azure", "Docker"], False),
    ("2018-2022: Software Engineer - Cyren", [
        "Designed and developed core components of email and network security products, improving threat detection quality.",
        "Maintained and evolved a complex SaaS platform with integrations across multiple internal and third-party services.",
        "Built phishing and spam detection engines used in production security pipelines.",
        "Drove technical POCs including headless browser analysis, image-recognition services, fast Hamming-distance processing, and locality-sensitive hashing."
    ], ['Go', 'C++', "Python", "ElasticSearch", "Redis", "Apache Kafka", "Prometheus", "Grafana", "Kibana", "Megalog", "Jaeger", "SQL", "k8s", "AWS", "Docker", "AI"], True),
    ("2016-2018: Software Engineer - Secdo", [
        "Defined and implemented Linux kernel probes with SystemTap and eBPF to monitor user-space activity with minimal overhead.",
        "Built a high-throughput pipeline that analyzed millions of system events per second on multicore environments.",
        "Implemented Windows kernel and driver-side collection components with user-space integration for efficient telemetry ingestion.",
        "Optimized detection data-processing workflows and developed behavioral models with the security research team."
    ], ['Python', 'C/C++', "Linux Kernel", "SystemTap", "eBPF", "Windows kernel", "Vertica", "SQL"], False),
    ("2012-2016: Software Engineer - Megabridge", [
        "Developed HLS video monitoring/control systems, including hardware verification and low-level embedded firmware.",
        "Designed device drivers and major firmware components for the BDSL product line, including Java-based management interfaces.",
        "Led hardware bring-up for multiple platforms and customized Linux root filesystems with Yocto.",
        "Established automated build infrastructure and CI processes for faster, repeatable delivery."
    ], ['C/C++', 'Python', "Linux kernel", "U-boot"], False),
    ("2009-2012: Software Engineer - Texas Instruments", [
        "Developed firmware for 802.11 ASIC programs, including pre- and post-silicon verification of WLAN devices.",
        "Contributed to WLAN Linux kernel development and built utilities for firmware build/debug workflows.",
        "Led migration from ClearCase to Git, improving engineering productivity by 30%.",
        "Built an internal wiki-based search system that improved discoverability of engineering knowledge."
    ], ['C/C++', 'Python', "Linux kernel", "ASIC", "Bare metal"], False),
]

if LONG_VERSION:
    experience += [
        ("2005-2007: Software Engineer - Broadlight Ltd.", [
            "Contributed to the development of the world's first fully integrated GPON chip.",
            "Supported the hardware team in post-production verification and pre-fab processes.",
            "Prepared firmware for design verification, post-production chip verification, and a series of drivers."
        ], ['C', 'Verilog', "Linux kernel", "Specman", "Bare metal"], True),
        ("2001-2004: Software Engineer - Terayon Israel/USA", [
            "Led a team of 5 engineers in the design and development of DSLAM systems, INTERCARD communication BSP for vxWorks, and various hardware drivers.",
            "Delivered customer-tailored systems, some of which are still in operation.",
            "Conducted employee training programs."
        ], ['C', "Linux kernel", "Bare metal"], False),
        ("1995-2000: Software Engineer - TDSoft", [
            "Developed infrastructures and low-level drivers for embedded real-time systems in C, C++, and Java.",
            "Participated in several development projects, including a Windows management system for strategic customers such as Deutsche Telekom and ECI.",
            "Designed and implemented significant parts of a V5 Class 5 switch (LE) simulator with ISDN support and developed GUI management for the simulator."
        ], ['C/C++', "Java", "Bare metal", "RTKernel", "pSOS"], False)
    ]


for position, details, technologies, page_break in experience:
    doc.add_heading(position, level=2)
    for detail in details:
        doc.add_paragraph(detail, style='List Bullet')
    if ADD_TECHOLOGIES and technologies:
        doc.add_heading('Technologies & Programming Languages', level=3)
        doc.add_paragraph(', '.join(technologies))
    if page_break:
        doc.add_page_break()

doc.add_heading('Education', level=1)
doc.add_paragraph('B.A. in Business Administration - Netanya Academic College (2007-2009)')
doc.add_paragraph('Software Practical Engineer - Tel Aviv University School of Practical Engineering')
doc.add_paragraph('Physics Department - St. Petersburg State University (1987-1989)')

# Save the document
suffix = {False: "_rt", True: ""}[IS_BACKEND]
file_path = f"./Arkady_Miasnikov_CV_Updated{suffix}.docx"
doc.save(file_path)

pdf_path = file_path.replace('.docx', '.pdf')
try:
    pypandoc.convert_file(
        file_path,
        'pdf',
        outputfile=pdf_path,
        extra_args=[
           '-V', 'geometry:top=0.5in',
           '-V', 'geometry:bottom=1in',
           '-V', 'geometry:left=1in',
           '-V', 'geometry:right=1in'
        ]
    )
except RuntimeError as e:
    print(f"DOCX generated: {file_path}")
    print("PDF generation failed.")
    print("Install missing LaTeX packages and rerun:")
    print("  sudo apt-get update")
    print("  sudo apt-get install -y pandoc texlive-latex-base texlive-latex-recommended texlive-latex-extra")
    raise

file_path
